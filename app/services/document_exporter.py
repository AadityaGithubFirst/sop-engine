"""Markdown to styled Microsoft Word (.docx) conversion.

Deliberately implements a small, predictable Markdown subset rather than
pulling in a heavyweight converter: headings, paragraphs, GFM tables, ordered
and unordered lists, blockquotes, fenced code, horizontal rules, and inline
bold/italic/code. Anything unrecognised degrades to plain text rather than
raising, because losing a government document to a stray character is worse
than losing its formatting.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.config import settings

logger = logging.getLogger(__name__)

# Inline spans: `code`, **bold**, __bold__, *italic*, _italic_
_INLINE_PATTERN = re.compile(
    r"(`[^`]+`|\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)"
)
_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.*)$")
_ORDERED_PATTERN = re.compile(r"^\s*(\d+)[.)]\s+(.*)$")
_UNORDERED_PATTERN = re.compile(r"^\s*[-*+]\s+(.*)$")
_RULE_PATTERN = re.compile(r"^\s*([-*_])\1{2,}\s*$")
_TABLE_DIVIDER_PATTERN = re.compile(r"^\s*\|?[\s:|-]+\|[\s:|-]*$")

ACCENT = RGBColor(0x1F, 0x3B, 0x63)  # departmental navy


def _safe_stem(document_id: str) -> str:
    """Reduce an identifier to a filename that cannot escape the output dir.

    Strips any directory component, replaces unsafe characters, and collapses
    dot runs so `../../escape` can never resolve outside `output_dir`.
    """
    stem = Path(document_id).name
    stem = re.sub(r"[^A-Za-z0-9._-]", "_", stem)
    stem = re.sub(r"\.{2,}", "_", stem).strip("._-")
    return stem or "sop_document"


class DocumentExportError(RuntimeError):
    """Raised when the .docx artifact cannot be produced or read back."""


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------
def _add_inline_runs(paragraph, text: str) -> None:
    """Append `text` to `paragraph`, honouring inline Markdown markers."""
    for part in _INLINE_PATTERN.split(text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif (part.startswith("**") and part.endswith("**")) or (
            part.startswith("__") and part.endswith("__")
        ):
            paragraph.add_run(part[2:-2]).bold = True
        elif (
            len(part) > 2
            and part[0] in "*_"
            and part[-1] == part[0]
        ):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Table parsing
# ---------------------------------------------------------------------------
def _split_row(line: str) -> List[str]:
    """Split a Markdown table row into trimmed cell values."""
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _is_table_start(lines: List[str], index: int) -> bool:
    """True when `lines[index]` is a header row followed by a divider row."""
    if "|" not in lines[index]:
        return False
    if index + 1 >= len(lines):
        return False
    return bool(_TABLE_DIVIDER_PATTERN.match(lines[index + 1]))


def _consume_table(lines: List[str], index: int) -> Tuple[List[List[str]], int]:
    """Collect a full Markdown table starting at `index`.

    Returns the rows (header first) and the index of the first line after it.
    """
    header = _split_row(lines[index])
    cursor = index + 2  # skip header and divider
    rows: List[List[str]] = [header]
    while cursor < len(lines) and "|" in lines[cursor] and lines[cursor].strip():
        cells = _split_row(lines[cursor])
        # Normalise ragged rows so python-docx never indexes out of range.
        if len(cells) < len(header):
            cells += [""] * (len(header) - len(cells))
        rows.append(cells[: len(header)])
        cursor += 1
    return rows, cursor


def _render_table(document: Document, rows: List[List[str]]) -> None:
    """Write a parsed Markdown table into the document with header shading."""
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_inline_runs(paragraph, value)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = ACCENT
    document.add_paragraph()


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------
def _apply_base_styles(document: Document) -> None:
    """Set a readable, print-friendly default style for the whole document."""
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)


def _add_heading(document: Document, level: int, text: str) -> None:
    heading = document.add_heading(level=min(level, 4))
    _add_inline_runs(heading, text)
    for run in heading.runs:
        run.font.color.rgb = ACCENT


def _add_code_block(document: Document, code_lines: List[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def markdown_to_document(markdown: str, title: Optional[str] = None) -> Document:
    """Convert a Markdown string into an in-memory `python-docx` Document."""
    document = Document()
    _apply_base_styles(document)

    if title:
        heading = document.add_heading(title, level=0)
        for run in heading.runs:
            run.font.color.rgb = ACCENT
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    lines = markdown.replace("\r\n", "\n").split("\n")
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            index += 1
            code: List[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            _add_code_block(document, code)
            index += 1  # closing fence
            continue

        if _RULE_PATTERN.match(stripped):
            rule = document.add_paragraph()
            rule.add_run("_" * 68).font.color.rgb = ACCENT
            index += 1
            continue

        if _is_table_start(lines, index):
            rows, index = _consume_table(lines, index)
            _render_table(document, rows)
            continue

        heading_match = _HEADING_PATTERN.match(stripped)
        if heading_match:
            _add_heading(document, len(heading_match.group(1)), heading_match.group(2).strip())
            index += 1
            continue

        if stripped.startswith(">"):
            quote = document.add_paragraph()
            quote.paragraph_format.left_indent = Pt(24)
            _add_inline_runs(quote, stripped.lstrip("> ").strip())
            for run in quote.runs:
                run.italic = True
            index += 1
            continue

        ordered = _ORDERED_PATTERN.match(line)
        if ordered:
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_runs(paragraph, ordered.group(2).strip())
            index += 1
            continue

        unordered = _UNORDERED_PATTERN.match(line)
        if unordered:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_runs(paragraph, unordered.group(1).strip())
            index += 1
            continue

        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, stripped)
        index += 1

    return document


def export_markdown_to_docx(
    markdown: str,
    document_id: str,
    output_dir: Optional[Path] = None,
    title: Optional[str] = None,
) -> Path:
    """Render `markdown` to a .docx file named after `document_id`.

    The file is written under `settings.OUTPUT_DIR` (`/tmp/sop_engine` on
    POSIX). Returns the path to the written file.
    """
    target_dir = Path(output_dir) if output_dir else settings.output_path
    target_dir.mkdir(parents=True, exist_ok=True)

    destination = target_dir / f"{_safe_stem(document_id)}.docx"

    try:
        document = markdown_to_document(markdown, title=title)
        document.save(str(destination))
    except Exception as exc:  # noqa: BLE001 - surfaced as a typed error
        raise DocumentExportError(f"Failed to export {document_id} to .docx: {exc}") from exc

    if not destination.exists() or destination.stat().st_size == 0:
        raise DocumentExportError(f"Exported file {destination} is missing or empty.")

    logger.info("Exported %s (%d bytes)", destination, destination.stat().st_size)
    return destination


def resolve_docx_path(document_id: str, output_dir: Optional[Path] = None) -> Optional[Path]:
    """Return the stored .docx for `document_id`, or None if it is absent.

    The identifier is sanitised the same way it was on write, which also keeps
    path traversal (`../`) out of the lookup.
    """
    target_dir = Path(output_dir) if output_dir else settings.output_path
    candidate = target_dir / f"{_safe_stem(document_id)}.docx"
    return candidate if candidate.is_file() else None
